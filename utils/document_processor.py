import os
import logging
import unicodedata
from PyPDF2 import PdfReader
from werkzeug.utils import secure_filename
from markitdown import MarkItDown
from openai import OpenAI
from docx import Document as DocxDocument
import datetime

# Set up logging with more detailed format
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize OpenAI client for MarkItDown
try:
    openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    md = MarkItDown(llm_client=openai_client, llm_model="gpt-4o")
    logger.debug("Successfully initialized MarkItDown with OpenAI integration")
except Exception as e:
    logger.error(f"Failed to initialize MarkItDown: {str(e)}", exc_info=True)
    raise

def normalize_filename(filename):
    """
    Normalize Unicode filename while preserving Chinese characters.
    """
    try:
        # Normalize Unicode characters (NFC form is preferred for most systems)
        normalized = unicodedata.normalize('NFC', filename)

        # Replace problematic characters but keep Chinese characters
        safe_chars = []
        for char in normalized:
            if unicodedata.category(char).startswith(('Lu', 'Ll', 'Nd', 'Han')):  # Letters, numbers, and Chinese characters
                safe_chars.append(char)
            elif char in ('-', '_', '.'):  # Allow certain punctuation
                safe_chars.append(char)
            else:
                safe_chars.append('_')  # Replace other characters with underscore

        return ''.join(safe_chars)
    except Exception as e:
        logger.error(f"Error normalizing filename: {str(e)}")
        return secure_filename(filename)  # Fallback to secure_filename

def extract_doc_metadata(file_path):
    """Extract metadata from a Word document."""
    try:
        doc = DocxDocument(file_path)
        core_properties = doc.core_properties
        return {
            "author": core_properties.author,
            "created": core_properties.created.isoformat() if core_properties.created else None,
            "modified": core_properties.modified.isoformat() if core_properties.modified else None,
            "title": core_properties.title,
            "subject": core_properties.subject,
            "keywords": core_properties.keywords,
            "category": core_properties.category,
            "paragraphs": len(doc.paragraphs),
            "sections": len(doc.sections)
        }
    except Exception as e:
        logger.warning(f"Could not extract metadata: {str(e)}")
        return {}

def extract_text_using_python_docx(file_path):
    """Fallback method to extract text using python-docx."""
    try:
        logger.debug("Attempting text extraction using python-docx")
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        if not text.strip():
            raise ValueError("No text content found in document")
        logger.info("Successfully extracted text using python-docx")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text using python-docx: {str(e)}")
        raise

def extract_text_from_word(file_path):
    """Extract text content from a Word document with fallback mechanisms."""
    logger.debug(f"Attempting to extract text from Word document: {file_path}")

    errors = []
    metadata = {}

    try:
        # First try to extract metadata
        try:
            metadata = extract_doc_metadata(file_path)
            logger.debug(f"Extracted metadata: {metadata}")
        except Exception as e:
            logger.warning(f"Metadata extraction failed: {str(e)}")
            errors.append(("metadata", str(e)))

        # Verify file exists and is readable
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"Word document not found at path: {file_path}")

        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"No read permission for file: {file_path}")

        # Check file size and log it
        file_size = os.path.getsize(file_path)
        logger.debug(f"File size: {file_size / 1024:.1f}KB")

        if file_size == 0:
            raise ValueError("File is empty (0 bytes)")

        # First attempt: Try MarkItDown
        try:
            logger.debug("Attempting text extraction using MarkItDown")
            result = md.convert(file_path)
            if result and hasattr(result, 'text_content') and result.text_content.strip():
                logger.info("Successfully extracted text using MarkItDown")
                return result.text_content, metadata
            else:
                raise ValueError("MarkItDown returned empty content")
        except Exception as e:
            logger.warning(f"MarkItDown extraction failed: {str(e)}")
            errors.append(("markitdown", str(e)))

            # Second attempt: Try python-docx
            try:
                text = extract_text_using_python_docx(file_path)
                if text.strip():
                    logger.info("Successfully extracted text using fallback method (python-docx)")
                    return text, metadata
                else:
                    raise ValueError("python-docx returned empty content")
            except Exception as e:
                logger.error(f"All extraction methods failed for {file_path}")
                errors.append(("python-docx", str(e)))
                raise ValueError(f"Failed to extract text using all available methods. Errors: {errors}")

    except Exception as e:
        logger.error(f"Failed to process Word document: {str(e)}", exc_info=True)
        # Clean up any temporary files
        try:
            temp_dir = os.path.join(os.path.dirname(file_path), '.markitdown_temp')
            if os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir)
                logger.debug(f"Cleaned up temporary directory: {temp_dir}")
        except Exception as cleanup_error:
            logger.error(f"Error cleaning up temporary files: {str(cleanup_error)}")
        raise

def extract_text_from_pdf(file_path):
    """Extract text content from a PDF file."""
    logger.debug(f"Attempting to extract text from PDF: {file_path}")
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"Cannot read PDF file: {file_path}")

        # Check file size
        check_file_size(file_path)

        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                logger.debug(f"Processing page {page_num + 1}")
                text += page.extract_text()

        if not text.strip():
            raise ValueError("No text content extracted from PDF")

        logger.debug("PDF text extraction successful")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {str(e)}", exc_info=True)
        raise

def check_file_size(file_path, max_size_mb=16):
    """Check if file size is within acceptable limits."""
    try:
        file_size = os.path.getsize(file_path)
        max_size_bytes = max_size_mb * 1024 * 1024

        logger.debug(f"Checking file size for {file_path}: {file_size / 1024 / 1024:.2f}MB")

        if file_size > max_size_bytes:
            raise ValueError(f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds maximum allowed size ({max_size_mb}MB)")
        if file_size == 0:
            raise ValueError("File is empty")
        return file_size
    except OSError as e:
        raise ValueError(f"Error checking file size: {str(e)}")

def allowed_file(filename):
    """Check if the file extension is allowed."""
    ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
    try:
        if '.' not in filename:
            logger.error("No file extension found")
            return False

        extension = filename.rsplit('.', 1)[1].lower()
        if not extension in ALLOWED_EXTENSIONS:
            logger.error(f"Extension {extension} not in allowed extensions")
            return False

        if not os.path.splitext(filename)[0]:
            logger.error("No filename part found")
            return False

        return True
    except Exception as e:
        logger.error(f"Error checking file extension: {str(e)}")
        return False

def process_document(file_path):
    """Process a document file and extract its text content with metadata."""
    logger.debug(f"Starting document processing for: {file_path}")
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()
        logger.debug(f"Detected file extension: {file_extension}")

        if not file_extension:
            logger.error("No file extension detected")
            raise ValueError("Unable to determine file type - no extension found")

        # Verify file is not empty
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.error("Empty file detected")
            raise ValueError("The uploaded file is empty")

        logger.debug(f"File size: {file_size / 1024:.1f}KB")

        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path), {}
        elif file_extension in ['.doc', '.docx']:
            return extract_text_from_word(file_path)
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            raise ValueError(f"Unsupported file type: {file_extension}")

    except Exception as e:
        logger.error(f"Error in process_document: {str(e)}", exc_info=True)
        # Clean up the file if there was an error
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.debug(f"Cleaned up file {file_path} after error")
        except Exception as cleanup_error:
            logger.error(f"Error cleaning up file: {str(cleanup_error)}")
        raise